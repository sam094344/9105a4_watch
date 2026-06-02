import docx

doc = docx.Document("介绍页.docx")
print("=== Paragraphs inside 介绍页.docx ===")
for p in doc.paragraphs:
    if p.text.strip():
        print(p.text.strip())

print("=== Tables inside 介绍页.docx ===")
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            print(cell.text.strip())
